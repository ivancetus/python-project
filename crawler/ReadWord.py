# pip install python-docx
import re

import docx

path = 'E:\server\wordpress\portfolio\MSword\\test2.docx'

file=docx.Document(path)
# for i in file.paragraphs:
#     print(i.text.strip())

for i in file.paragraphs:
    if re.search("[A-Z][a-z]+.+[.,?:]", i.text):
        print(i.text)
    else:
        pass
        #print(i.text)
